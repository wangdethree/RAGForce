from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from src.schemas.ingestion import ParsedDocument, ImageBlock


class DocumentParser:
    """解析 PDF 和 Word 文档为结构化内容"""

    async def parse(self, file_path: str, file_type: str) -> ParsedDocument:
        if file_type == "pdf":
            return await self._parse_pdf(file_path)
        elif file_type in ("docx", "doc"):
            return await self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

    async def _parse_pdf(self, file_path: str) -> ParsedDocument:
        doc = fitz.open(file_path)
        pages = []
        images = []

        for page_num, page in enumerate(doc):
            text = page.get_text()
            pages.append({"page_num": page_num + 1, "text": text, "tables": []})

            for img_index, img in enumerate(page.get_images(full=True)):
                xref = img[0]
                base_image = doc.extract_image(xref)
                images.append(
                    ImageBlock(
                        page_num=page_num + 1,
                        index=img_index,
                        image_bytes=base_image["image"],
                        ext=base_image["ext"],
                    )
                )

        full_text = "\n\n".join(p["text"] for p in pages)
        return ParsedDocument(
            text=full_text,
            pages=pages,
            images=images,
            tables=[],
        )

    async def _parse_docx(self, file_path: str) -> ParsedDocument:
        doc = DocxDocument(file_path)
        pages = []
        images = []
        full_text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text_parts.append(para.text)

        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.reltype:
                images.append(
                    ImageBlock(
                        page_num=1,
                        index=i,
                        image_bytes=rel.target_part.blob,
                        ext=rel.target_part.partname.split(".")[-1],
                    )
                )

        return ParsedDocument(
            text="\n".join(full_text_parts),
            pages=pages,
            images=images,
            tables=[],
        )


document_parser = DocumentParser()
